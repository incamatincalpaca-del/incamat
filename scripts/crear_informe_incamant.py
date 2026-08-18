from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


OUT = Path(r"C:\IncaMant\outputs\Informe_EFSRT_INCAMAT_Alyson_Diaz_V6_Gantt_Visual.docx")
GANTT_IMAGE = Path(r"C:\IncaMant\outputs\Cronograma_Gantt_INCAMAT.png")

NAVY = "000000"
BLUE = "000000"
LIGHT_BLUE = "D9D9D9"
GRAY = "404040"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_cm):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Página ")
    run.font.size = Pt(9)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def add_toc(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
    paragraph._p.append(fld)


def enable_field_update(doc):
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def style_run(run, size=10, bold=False, color=None, italic=False):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_text(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        style_run(r, 10, True)
        r = p.add_run(text[len(bold_prefix):])
        style_run(r)
    else:
        r = p.add_run(text)
        style_run(r)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text)
    style_run(r)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    style_run(r, 10, True, NAVY)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    style_run(r, 10, italic=True, color=GRAY)
    return p


def add_figure(doc, image_path, caption, explanation):
    path = Path(image_path)
    if not path.exists():
        return
    new_page(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(15.5))
    add_caption(doc, caption)
    add_text(doc, explanation)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header = table.rows[0]
    set_repeat_table_header(header)
    for i, value in enumerate(headers):
        cell = header.cells[i]
        if widths:
            set_cell_width(cell, widths[i])
        set_cell_shading(cell, NAVY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(value)
        style_run(r, 10, True, "000000")
    for row_values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_values):
            if widths:
                set_cell_width(cells[i], widths[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(str(value))
            style_run(r, 10)
    for row in table.rows:
        for cell in row.cells:
            cell.margin_top = Cm(0.08)
            cell.margin_bottom = Cm(0.08)
            cell.margin_left = Cm(0.1)
            cell.margin_right = Cm(0.1)
    return table


def add_gantt_table(doc):
    headers = ["Actividad del proyecto INCAMAT", "Inicio real", "Fin real", "P1", "P2", "P3", "P4", "P5", "P6", "P7", "P8", "P9", "P10"]
    activities = [
        ("Levantamiento de requerimientos y revisión de archivos Excel", [1]),
        ("Definición de estructura: planta, localizaciones, áreas y máquinas", [1, 2]),
        ("Diseño de base de datos y configuración Docker/MariaDB", [2]),
        ("Desarrollo del módulo de localizaciones y máquinas", [3]),
        ("Implementación de ficha técnica y código QR por máquina", [4]),
        ("Desarrollo de incidencias y órdenes de mantenimiento", [5, 6]),
        ("Implementación de importaciones, validación y exportación CSV", [6]),
        ("Clasificación de repuestos por área, familia y criticidad", [7, 8]),
        ("Solicitudes, movimientos de stock y control de repuestos", [8]),
        ("Dashboard e indicadores de mantenimiento y repuestos", [9]),
        ("Pruebas funcionales, correcciones y mejora de interfaz", [9, 10]),
        ("Elaboración de informe, anexos y evidencias", [10]),
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header = table.rows[0]
    set_repeat_table_header(header)
    widths = [5.4, 1.65, 1.65] + [0.62] * 10
    for i, value in enumerate(headers):
        cell = header.cells[i]
        set_cell_width(cell, widths[i])
        set_cell_shading(cell, "D9D9D9")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(value)
        style_run(r, 8, True)
    for activity, active_weeks in activities:
        cells = table.add_row().cells
        for i, cell in enumerate(cells):
            set_cell_width(cell, widths[i])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
            if i == 0:
                value = activity
            elif i in (1, 2):
                value = "____/____/____"
            else:
                week = i - 2
                value = "X" if week in active_weeks else ""
                if week in active_weeks:
                    set_cell_shading(cell, "A6A6A6")
            r = p.add_run(value)
            style_run(r, 8, bold=(i >= 3 and value == "X"))
    return table


def build_gantt_image():
    activities = [
        ("Levantamiento de requerimientos y revisión de archivos Excel", [1]),
        ("Estructura de planta, localizaciones, áreas y máquinas", [1, 2]),
        ("Diseño de base de datos y configuración Docker/MariaDB", [2]),
        ("Módulo de localizaciones y máquinas", [3]),
        ("Ficha técnica y código QR por máquina", [4]),
        ("Incidencias y órdenes de mantenimiento", [5, 6]),
        ("Importaciones, validación y exportación CSV", [6]),
        ("Repuestos por área, familia y criticidad", [7, 8]),
        ("Solicitudes, movimientos y control de stock", [8]),
        ("Dashboard e indicadores", [9]),
        ("Pruebas funcionales y mejora de interfaz", [9, 10]),
        ("Informe, anexos y evidencias", [10]),
    ]
    width, height = 1800, 1190
    left, top, header_h, row_h, period_w = 800, 175, 80, 66, 92
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    try:
        title_font = ImageFont.truetype("arialbd.ttf", 34)
        header_font = ImageFont.truetype("arialbd.ttf", 19)
        body_font = ImageFont.truetype("arial.ttf", 18)
        small_font = ImageFont.truetype("arial.ttf", 16)
    except OSError:
        title_font = ImageFont.load_default()
        header_font = ImageFont.load_default()
        body_font = ImageFont.load_default()
        small_font = ImageFont.load_default()
    navy, blue, gray, grid = "#17365D", "#2F75B5", "#F2F2F2", "#BFBFBF"
    draw.text((70, 52), "CRONOGRAMA DE GANTT - PROYECTO INCAMAT", fill=navy, font=title_font)
    draw.text((70, 108), "P1-P10 representan períodos de trabajo. Completar las fechas reales en el informe.", fill="#404040", font=small_font)
    draw.rectangle((70, top, left, top + header_h), fill=navy)
    draw.text((95, top + 27), "Actividades del proyecto", fill="white", font=header_font)
    for index in range(10):
        x0 = left + index * period_w
        draw.rectangle((x0, top, x0 + period_w, top + header_h), fill=navy)
        label = f"P{index + 1}"
        box = draw.textbbox((0, 0), label, font=header_font)
        draw.text((x0 + (period_w - (box[2] - box[0])) / 2, top + 27), label, fill="white", font=header_font)
    for row, (activity, periods) in enumerate(activities):
        y0 = top + header_h + row * row_h
        fill = "#FFFFFF" if row % 2 == 0 else gray
        draw.rectangle((70, y0, left + 10 * period_w, y0 + row_h), fill=fill)
        draw.text((92, y0 + 21), activity, fill="#1F1F1F", font=body_font)
        for col in range(11):
            x = left + col * period_w
            draw.line((x, y0, x, y0 + row_h), fill=grid, width=2)
        draw.line((70, y0 + row_h, left + 10 * period_w, y0 + row_h), fill=grid, width=2)
        for period in periods:
            x0 = left + (period - 1) * period_w + 11
            x1 = left + period * period_w - 11
            draw.rounded_rectangle((x0, y0 + 16, x1, y0 + row_h - 16), radius=9, fill=blue)
    draw.rectangle((70, top, left + 10 * period_w, top + header_h + len(activities) * row_h), outline=navy, width=2)
    draw.text((70, height - 68), "Fuente: elaboración propia, basada en las actividades desarrolladas en INCAMAT.", fill="#404040", font=small_font)
    image.save(GANTT_IMAGE)
    return GANTT_IMAGE


def add_gantt_figure(doc):
    image_path = build_gantt_image()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Cm(16.0))


def new_page(doc):
    doc.add_page_break()


def make_document():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.2)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for section in doc.sections:
        footer = section.footer
        add_page_number(footer.paragraphs[0])

    # Cover
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("INFORME DE EXPERIENCIA FORMATIVA\nEN SITUACIÓN REAL DE TRABAJO")
    style_run(r, 12, True, NAVY)
    p.paragraph_format.space_after = Pt(24)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MÓDULO II")
    style_run(r, 12, True)
    p.paragraph_format.space_after = Pt(26)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DISEÑO E IMPLEMENTACIÓN DE INCAMAT: SISTEMA WEB PARA LA GESTIÓN DE MANTENIMIENTO, INCIDENCIAS Y REPUESTOS")
    style_run(r, 12, True, NAVY)
    p.paragraph_format.space_after = Pt(28)

    cover_data = [
        ("EMPRESA:", "INCALPACA TPX - AREQUIPA"),
        ("ESTUDIANTE:", "ALYSON NOELY DIAZ MAMANI"),
        ("ESPECIALIDAD:", "[CONFIRMAR ESPECIALIDAD DE ALYSON]"),
        ("ASESOR:", "ARES SILVA"),
        ("ÁREA ASIGNADA:", "MANTENIMIENTO"),
        ("SUPERVISOR DE EMPRESA:", "[CONFIRMAR NOMBRE DEL SUPERVISOR]"),
    ]
    table = doc.add_table(rows=len(cover_data), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row, (label, value) in zip(table.rows, cover_data):
        set_cell_width(row.cells[0], 5.5)
        set_cell_width(row.cells[1], 8.5)
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        for i, text in enumerate((label, value)):
            p = row.cells[i].paragraphs[0]
            r = p.add_run(text)
            style_run(r, 12, bold=(i == 0), color=NAVY if i == 0 else None)
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AREQUIPA - PERÚ\nAGOSTO 2026")
    style_run(r, 12, True)

    # Index
    new_page(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ÍNDICE")
    style_run(r, 10, True, NAVY)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    add_toc(p)
    new_page(doc)
    add_heading(doc, "1. Introducción")
    add_text(doc, "El presente informe describe la experiencia formativa desarrollada mediante el diseño e implementación de INCAMAT, un sistema web orientado a la gestión de mantenimiento industrial, incidencias, máquinas y repuestos técnicos para la planta Incalpaca TPX en Arequipa.")
    add_text(doc, "El proyecto surgió ante la necesidad de ordenar información operativa que se encontraba distribuida en archivos Excel, tales como localizaciones de planta, máquinas, solicitudes de mantenimiento y catálogos de repuestos. Esta situación dificultaba consultar el historial de una máquina, conocer qué área solicita más repuestos, controlar el stock y priorizar las intervenciones de mantenimiento.")
    add_text(doc, "INCAMAT integra estos procesos en una sola plataforma. La solución permite importar información desde Excel, organizar máquinas por áreas y localizaciones, identificar equipos por código QR, registrar incidencias con evidencias fotográficas, atender órdenes de mantenimiento y controlar la salida de repuestos. De esta manera, se mejora la trazabilidad de las actividades y se facilita la toma de decisiones del área técnica y administrativa.")

    add_heading(doc, "2. Objetivo del proyecto")
    add_text(doc, "Diseñar e implementar un sistema web de gestión de mantenimiento e inventario técnico para Incalpaca TPX, utilizando una arquitectura basada en React, Node.js, Docker y MariaDB, con el propósito de centralizar la información de máquinas, incidencias, órdenes de trabajo y repuestos por área.")
    add_heading(doc, "2.1 Objetivos específicos", 2)
    for text in [
        "Organizar la estructura de planta mediante niveles de localización, áreas, subáreas y máquinas.",
        "Implementar la importación validada de archivos Excel para máquinas, áreas, repuestos y solicitudes de mantenimiento.",
        "Registrar incidencias con prioridad, fecha, descripción y evidencia fotográfica mediante códigos QR por máquina.",
        "Gestionar órdenes preventivas, correctivas, predictivas, proactivas y autónomas con horas de inicio y cierre, diagnóstico, trabajo realizado y repuestos utilizados.",
        "Clasificar repuestos en familias mecánica, eléctrica y electrónica, e incorporar control de stock, criticidad y solicitudes por área.",
        "Presentar indicadores de mantenimiento y consumo de repuestos para apoyar al supervisor y al área administrativa.",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "3. Ejecución y control del proyecto")
    add_heading(doc, "3.1 Perfil de la empresa", 2)
    add_text(doc, "Incalpaca TPX es una empresa industrial ubicada en Arequipa, Perú, cuya operación comprende procesos textiles y de confección. En su planta intervienen áreas como Hilandería, Acabado de Telas, Tejido Plano, Tejido Punto, CTP, Calidad, Mantenimiento, Almacenes, Tintorería y Zurcido. La diversidad de equipos y repuestos requiere una gestión ordenada de activos, mantenimientos y materiales técnicos.")
    add_text(doc, "De acuerdo con el organigrama institucional proporcionado por la empresa, el área de Mantenimiento pertenece a la Gerencia de Operaciones. La experiencia formativa y el proyecto INCAMAT se desarrollan desde esta área, por lo que la solución prioriza el registro de incidencias, órdenes de trabajo, disponibilidad de máquinas y repuestos técnicos.")
    add_heading(doc, "Misión y visión institucional", 3)
    add_text(doc, "Misión: “Creamos productos únicos a partir de Alpaca y Vicuña, siempre ligados a historias, conceptos innovadores y utilizando procesos sustentables e integrados. A través de ellos, nos asociamos con empresas, marcas y diseñadores exigentes y sofisticados. Buscamos trascender en el tiempo para asegurar la continuidad de nuestros grupos de interés, especialmente los productores de estas fibras nobles” (Incalpaca, s.f.).")
    add_text(doc, "Visión: “Ser líder mundial en la confección y comercialización de prendas únicas de Alpaca y Vicuña, conectadas con diversas comunidades, dentro y fuera del país, que nos ayuden a promover una moda más justa, sustentable y verdaderamente peruana” (Incalpaca, s.f.).")
    add_text(doc, "[Completar únicamente el número de colaboradores si la empresa autoriza incorporarlo en el informe.]", bold_prefix="[Completar")

    add_heading(doc, "3.1.1 Comunicación interna y recursos de trabajo", 3)
    add_text(doc, "En el área de Mantenimiento, la coordinación interna se realiza mediante la aplicación utilizada por la oficina para el registro y seguimiento de requerimientos, complementada con reuniones de coordinación. Este mecanismo permite comunicar incidencias, revisar prioridades, asignar actividades y dar seguimiento a las órdenes de trabajo.")
    add_text(doc, "Los principales recursos de trabajo relacionados con el proyecto fueron una computadora con acceso a la red local, archivos Excel históricos, la plataforma INCAMAT, Docker para ejecutar los servicios del sistema y las herramientas de consulta empleadas por el área. Estos recursos permitieron revisar la información existente, organizarla y validar el funcionamiento de los módulos desarrollados.")

    add_heading(doc, "3.2 Planificación y metodología de trabajo", 2)
    add_text(doc, "El desarrollo se organizó por etapas iterativas. En cada etapa se revisó la información disponible, se definieron las reglas de negocio, se implementó el módulo correspondiente y se validó el resultado en el entorno local con Docker. Esta metodología permitió ajustar el sistema conforme aparecían nuevos archivos Excel, necesidades operativas y observaciones de los usuarios.")
    add_table(doc,
        ["Etapa", "Actividades desarrolladas", "Entregable"],
        [
            ["1. Levantamiento", "Revisión de archivos Excel, estructura de planta, máquinas, repuestos y solicitudes.", "Modelo inicial de datos y reglas de importación."],
            ["2. Diseño", "Definición de módulos, navegación, campos, estados, prioridades y relaciones entre entidades.", "Prototipo funcional de INCAMAT."],
            ["3. Construcción", "Desarrollo de frontend, API, base de datos, importaciones y contenedores Docker.", "Módulos operativos integrados."],
            ["4. Validación", "Pruebas de carga, filtros, reportes, QR, órdenes, movimientos y exportación CSV.", "Correcciones y mejoras de usabilidad."],
            ["5. Documentación", "Registro de decisiones, evidencias, resultados y recomendaciones de continuidad.", "Informe de experiencia formativa."],
        ], [2.6, 9.0, 3.4])
    add_caption(doc, "Tabla 1. Etapas de desarrollo del proyecto. Fuente: elaboración propia.")

    add_heading(doc, "3.2.1 Métodos y criterios utilizados en el proyecto", 3)
    add_text(doc, "El proyecto se alinea de manera referencial con el enfoque de gestión de activos de ISO 55000, debido a que INCAMAT organiza información de activos físicos, riesgos, desempeño, mantenimiento y decisiones sobre repuestos. Esta referencia no implica que la empresa se encuentre certificada bajo dicha norma; se utiliza únicamente como marco conceptual para ordenar el alcance del sistema (ISO, 2024).")
    add_text(doc, "Como método incorporado en la aplicación, se definió una matriz de criticidad de repuestos. La evaluación considera cuatro factores: impacto en la producción, tiempo de reposición, disponibilidad de alternativa e impacto económico. El resultado permite registrar una criticidad baja, media, alta o crítica; sin embargo, la calificación definitiva requiere validación técnica por el responsable de Mantenimiento o Almacén.")
    add_text(doc, "Los siguientes métodos se plantean como una etapa posterior y no se consideran resultados alcanzados en este informe: AMFE/FMEA para identificar modos de falla, efectos y prioridades de tratamiento; análisis de Pareto para priorizar fallas y repuestos con mayor frecuencia; y los indicadores MTTR y MTBF una vez que existan órdenes cerradas y horas de operación validadas. IEC 60812 describe el AMFE/FMEA como un método sistemático para identificar modos de falla y sus efectos (IEC, 2018).")

    add_heading(doc, "3.2.2 Cronograma de Gantt del proyecto INCAMAT", 3)
    add_text(doc, "El siguiente cronograma representa la secuencia de trabajo del proyecto INCAMAT, desde el análisis de la información entregada hasta las pruebas y la documentación. Cada período (P1 a P10) corresponde a una fase de trabajo y conserva el orden real en que se desarrollaron los módulos. Las columnas de inicio y fin se dejan en blanco para registrar las fechas reales de ejecución; no representan fechas estimadas.")
    add_gantt_figure(doc)
    add_caption(doc, "Figura 1. Cronograma de Gantt del proyecto INCAMAT. Fuente: elaboración propia, basada en las actividades desarrolladas en el proyecto.")

    add_heading(doc, "3.3 Problema identificado", 2)
    add_text(doc, "Durante la revisión de la información disponible se identificó que los registros de máquinas, áreas, localizaciones, repuestos y solicitudes se manejaban principalmente en archivos Excel independientes. Esto generaba riesgo de duplicados, dificultades para relacionar un repuesto con su área, poca visibilidad del historial de mantenimiento y demora al ubicar el equipo que presenta una falla.")
    add_text(doc, "Asimismo, la falta de un flujo único para reportar una incidencia y cerrarla con diagnóstico, evidencia, tiempo de atención y repuestos utilizados limitaba el seguimiento de las órdenes de trabajo y el cálculo de indicadores técnicos.")

    add_heading(doc, "3.4 Desarrollo de la pasantía", 2)
    add_text(doc, "Se desarrolló INCAMAT como una aplicación web de uso interno. La solución se organizó en módulos funcionales para asegurar que cada usuario consulte y registre la información que corresponde a su actividad.")
    add_table(doc,
        ["Módulo", "Función implementada", "Beneficio"],
        [
            ["Dashboard", "Muestra equipos, alertas de stock, órdenes, incidencias e indicadores por área.", "Visión rápida para supervisión."],
            ["Localizaciones y máquinas", "Organiza la planta por niveles y asigna cada máquina a su área; incluye ficha técnica y QR.", "Ubicación y consulta rápida de activos."],
            ["Incidencias", "El operario reporta una falla por QR con prioridad, fecha, descripción y fotografía.", "Comunicación directa con mantenimiento."],
            ["Mantenimientos", "Administra órdenes preventivas, correctivas, predictivas, proactivas y autónomas.", "Trazabilidad de inicio, atención y cierre."],
            ["Repuestos", "Clasifica repuestos por familia técnica, área, criticidad, stock y movimientos.", "Mejor control de inventario técnico."],
            ["Importaciones", "Valida Excel, detecta duplicados y conserva el historial de carga.", "Migración segura de información existente."],
        ], [3.2, 7.2, 4.6])
    add_caption(doc, "Tabla 3. Módulos principales de INCAMAT. Fuente: elaboración propia.")

    add_heading(doc, "3.4.1 Desarrollo de la solución INCAMAT", 3)
    add_text(doc, "El alcance del proyecto comprende el control interno de información de mantenimiento e inventario técnico. INCAMAT no reemplaza un sistema ERP ni calcula todavía indicadores de producción como OEE con valores reales; deja preparada la estructura para integrarlos cuando el área proporcione horas programadas, horas de operación, unidades producidas y registros de calidad validados.")
    add_table(doc,
        ["Rol", "Acciones principales dentro del sistema"],
        [
            ["Operario", "Uso previsto: escanear QR y registrar una incidencia con prioridad, fecha, descripción y evidencia."],
            ["Técnico", "Uso previsto: iniciar y atender órdenes; registrar diagnóstico, trabajo, tiempo, repuestos y prueba final."],
            ["Supervisor", "Uso previsto: priorizar órdenes, revisar incidencias, validar cierres y consultar indicadores por área."],
            ["Almacén / administrativo", "Uso previsto: verificar stock, registrar movimientos, atender solicitudes y revisar consumo por área."],
            ["Administrador", "Uso previsto: gestionar importaciones, catálogos y estructura de planta."],
        ], [3.6, 11.4])
    add_caption(doc, "Tabla 4. Alcance por perfil de usuario. Fuente: elaboración propia.")

    add_heading(doc, "3.4.2 Arquitectura tecnológica", 3)
    add_text(doc, "La interfaz fue desarrollada con React y Vite. El servidor de aplicación se implementó con Node.js y Express, mientras que MariaDB almacena la información estructurada. Docker permite ejecutar los servicios de frontend, backend y base de datos de forma integrada y reproducible en un entorno local.")
    add_table(doc,
        ["Capa", "Tecnología", "Responsabilidad"],
        [
            ["Presentación", "React + Vite", "Interfaz dinámica para usuarios de planta, técnicos y supervisores."],
            ["Lógica", "Node.js + Express", "API, validaciones, importación y reglas del negocio."],
            ["Datos", "MariaDB", "Persistencia de máquinas, áreas, incidencias, órdenes, repuestos y movimientos."],
            ["Despliegue local", "Docker Compose", "Ejecución coordinada de servicios y base de datos."],
        ], [3.2, 4.0, 7.8])
    add_caption(doc, "Tabla 5. Arquitectura tecnológica del proyecto. Fuente: elaboración propia.")

    add_heading(doc, "3.4.3 Flujo de atención de incidencias", 3)
    for text in [
        "El operario escanea el código QR de la máquina y reporta la incidencia desde el formulario asociado al equipo.",
        "El reporte registra área, máquina, prioridad, fecha, descripción y evidencia inicial; se genera una orden correctiva pendiente.",
        "El técnico inicia la orden y el sistema registra la hora de inicio para medir el tiempo de atención.",
        "Durante la atención se registra el diagnóstico, tipo de causa, trabajo realizado, repuestos usados y fotografía final.",
        "Al cerrar la orden se guarda la hora final, el estado de prueba y se actualiza la condición operativa de la máquina según corresponda.",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "3.4.4 Gestión de repuestos", 3)
    add_text(doc, "Los repuestos se organizaron por área de uso y familia técnica: mecánica, eléctrica o electrónica. Para facilitar la priorización se implementó una evaluación de criticidad basada en impacto en producción, tiempo de reposición, disponibilidad de alternativa e impacto económico. La suma de estos factores permite clasificar cada repuesto como crítico, alto, medio o bajo.")
    add_text(doc, "También se incorporó la verificación de stock físico, ubicación de almacén, movimientos de ingreso o salida, asociación del consumo a una orden de mantenimiento y registro de solicitudes por área. Esta información permite identificar los repuestos más solicitados y la antigüedad de cada requerimiento.")

    add_heading(doc, "3.4.5 Modelo de datos e integridad de la información", 3)
    add_text(doc, "La información se estructuró mediante entidades relacionadas. Esta organización permite evitar que una máquina, una incidencia o un repuesto se registren de forma aislada. Cada registro conserva identificadores y relaciones para facilitar consultas por área, equipo, fecha o estado.")
    add_table(doc,
        ["Entidad", "Información principal", "Relación dentro de INCAMAT"],
        [
            ["Localizaciones", "Niveles de planta, división, zona y subzona.", "Agrupan las áreas y permiten ubicar activos."],
            ["Áreas", "Nombre, responsable, estado y ubicación jerárquica.", "Agrupan máquinas, incidencias, órdenes y repuestos de uso."],
            ["Máquinas", "Código, nombre, área, marca, modelo y estado.", "Se vinculan con QR, incidencias y mantenimientos."],
            ["Incidencias", "Prioridad, fecha, descripción, evidencia y estado.", "Generan o alimentan órdenes correctivas."],
            ["Mantenimientos", "Tipo, responsable, horas, diagnóstico, trabajo y prueba final.", "Registran la intervención realizada en una máquina."],
            ["Repuestos", "Código original, familia, criticidad, stock y ubicación.", "Se relacionan con áreas, solicitudes y movimientos."],
            ["Movimientos", "Ingreso, salida o ajuste de inventario.", "Permiten rastrear consumos y actualización de stock."],
        ], [3.0, 5.5, 6.5])
    add_caption(doc, "Tabla 6. Estructura de datos utilizada por INCAMAT. Fuente: elaboración propia.")

    add_heading(doc, "3.4.6 Proceso de importación desde Excel", 3)
    add_text(doc, "Debido a que los registros existentes se trabajan en Excel, se implementó un flujo de importación controlado. El usuario selecciona el catálogo que desea cargar, descarga una plantilla oficial, adjunta el archivo completado y revisa una vista previa antes de procesarlo. El sistema valida columnas obligatorias, muestra filas con errores y diferencia los registros que se pueden crear o actualizar.")
    for text in [
        "Las plantillas se mantienen separadas para áreas, máquinas, repuestos y mantenimientos provenientes de SRequest.",
        "Los códigos originales de los repuestos se conservan para no romper la trazabilidad con los archivos de origen.",
        "La importación mantiene un historial con archivo, tipo de carga, estado y opción de eliminar una carga errónea junto con sus datos asociados.",
        "La descarga en formato CSV permite revisar la información procesada nuevamente en Excel.",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "3.4.7 Diseño de la experiencia de usuario", 3)
    add_text(doc, "La interfaz se diseñó con una navegación lateral y módulos claramente identificados para reducir la curva de aprendizaje. Se utilizaron tarjetas, filtros, indicadores de estado y colores de prioridad para que el usuario identifique rápidamente información relevante, sin requerir conocimientos técnicos de programación.")
    add_text(doc, "En las órdenes se emplean estados diferenciados, tales como correctivos nuevos, en atención, esperando repuesto, pendiente de validación, preventivo planificado, autónomo, predictivo y proactivo. Esta clasificación permite visualizar el trabajo pendiente desde la perspectiva del responsable técnico y del supervisor.")

    add_heading(doc, "3.5 Pruebas y resultados obtenidos", 2)
    add_text(doc, "Como resultado se obtuvo una plataforma funcional que concentra la información antes dispersa en archivos de Excel. Se consolidaron las localizaciones de la planta, se organizaron las máquinas por área y se habilitó un catálogo de repuestos con relación a las unidades usuarias. El sistema mantiene los códigos originales de los catálogos importados para preservar la trazabilidad con los registros de la empresa.")
    add_text(doc, "La aplicación permite que los operarios reporten incidencias de forma guiada, mientras que el personal de mantenimiento dispone de una bandeja de órdenes con estados, prioridades, evidencias y tiempos de atención. Además, el dashboard presenta indicadores por área, lo cual aporta información para priorizar acciones y revisar el comportamiento del mantenimiento y del inventario técnico.")
    add_table(doc,
        ["Prueba realizada", "Resultado esperado", "Resultado"],
        [
            ["Importación de Excel", "Validar columnas y detectar registros duplicados antes de guardar.", "Cumplido: se muestra vista previa e historial de importación."],
            ["Reporte por QR", "Relacionar automáticamente la incidencia con la máquina consultada.", "Cumplido: el QR dirige al registro vinculado al equipo."],
            ["Atención de orden", "Guardar hora de inicio, diagnóstico, evidencia, repuestos y hora de cierre.", "Cumplido: la orden conserva la trazabilidad de la atención."],
            ["Movimiento de repuesto", "Descontar o ajustar stock y relacionar el movimiento con la actividad.", "Cumplido: se registra el movimiento y se actualiza el inventario."],
            ["Exportación CSV", "Descargar información operativa compatible con Excel.", "Cumplido: se habilitó descarga de reportes seleccionados."],
        ], [3.4, 6.7, 4.9])
    add_caption(doc, "Tabla 7. Matriz de pruebas funcionales. Fuente: elaboración propia.")

    add_heading(doc, "3.6 Indicadores e impacto esperado", 2)
    add_text(doc, "La plataforma habilita indicadores operativos para el seguimiento del mantenimiento y de los repuestos. Los valores reales deberán ser revisados periódicamente con el supervisor de cada área, pero el sistema deja disponible la información necesaria para calcularlos y compararlos por periodos.")
    add_table(doc,
        ["Indicador", "Forma de seguimiento", "Uso para la empresa"],
        [
            ["Órdenes pendientes y atrasadas", "Cantidad de órdenes según estado y fecha programada.", "Priorizar la atención diaria."],
            ["Tiempo de atención", "Diferencia entre hora de inicio y hora final de la orden.", "Revisar la duración de reparaciones."],
            ["MTTR", "Promedio de tiempo de reparación de fallas cerradas.", "Identificar oportunidades de mejora técnica."],
            ["MTBF", "Horas de operación entre fallas, con datos validados de operación.", "Analizar la confiabilidad de cada equipo."],
            ["Consumo por área", "Salidas de repuestos relacionadas con áreas y órdenes.", "Planificar solicitudes y stock técnico."],
            ["Repuestos críticos", "Nivel de criticidad, stock mínimo y verificación física.", "Evitar paradas por falta de material."],
        ], [3.1, 6.0, 5.9])
    add_caption(doc, "Tabla 8. Indicadores disponibles o preparados en INCAMAT. Fuente: elaboración propia.")

    add_heading(doc, "4. Análisis o balance crítico")
    add_heading(doc, "4.1 Tarea principal desarrollada durante la pasantía", 2)
    add_text(doc, "La tarea principal del proyecto fue el análisis, diseño e implementación de INCAMAT, una aplicación web de apoyo para el área de Mantenimiento. Las actividades documentadas del proyecto consistieron en revisar archivos Excel entregados por el área, identificar las relaciones entre áreas, máquinas, solicitudes y repuestos, y estructurar esta información en módulos dentro del sistema web.")
    add_text(doc, "Las actividades desarrolladas en INCAMAT incluyeron la organización de localizaciones y máquinas por área; la importación y validación de archivos Excel; el registro de incidencias con códigos QR y evidencias; el seguimiento de órdenes de mantenimiento; y el catálogo de repuestos con clasificación técnica, stock, criticidad y solicitudes por área. Esta descripción se basa en las funciones incorporadas en la versión actual de la aplicación.")
    add_heading(doc, "4.2 Dificultades encontradas y forma de abordarlas", 2)
    add_text(doc, "Las dificultades verificables del proyecto estuvieron relacionadas con la diversidad de los archivos Excel proporcionados: existen catálogos de máquinas, ubicaciones, solicitudes y repuestos con estructuras distintas. Además, fue necesario conservar los códigos originales de los repuestos y contar con validaciones para evitar duplicados durante las importaciones.")
    add_text(doc, "Como respuesta, INCAMAT incorpora plantillas de importación, vista previa, validaciones de columnas, detección de errores, filtros por área e historial de cargas. La clasificación de los repuestos se plantea de manera progresiva, ya que su validación técnica requiere la revisión del personal responsable del área.")
    add_heading(doc, "4.3 Logros alcanzados", 2)
    for text in [
        "Se transformó información de múltiples archivos Excel en una estructura única y consultable.",
        "Se estableció una relación entre áreas, máquinas, incidencias, mantenimientos y repuestos.",
        "Se diseñó una interfaz didáctica para que los usuarios de planta puedan consultar y registrar información con menos pasos.",
        "Se incorporó evidencia fotográfica, códigos QR y tiempos de atención para fortalecer la trazabilidad.",
        "Se habilitó la descarga de reportes CSV para conservar la compatibilidad con el trabajo administrativo en Excel.",
    ]:
        add_bullet(doc, text)
    add_heading(doc, "4.4 Dificultades y acciones de mejora", 2)
    add_text(doc, "La principal dificultad fue la heterogeneidad de los archivos de origen: nombres distintos para una misma área, columnas no uniformes, registros incompletos y repuestos sin clasificación previa. Para resolverlo se definieron plantillas de importación, validaciones de duplicados y una clasificación progresiva de repuestos por familia técnica.")
    add_text(doc, "También se identificó que algunos indicadores industriales, como OEE, MTBF y MTTR, requieren datos operativos adicionales: horas programadas de producción, horas efectivas de operación, producción total, productos conformes y tiempos de parada validados. INCAMAT ya registra una parte de esta información mediante las órdenes; la siguiente etapa consiste en integrar los datos de producción y validar las fórmulas con el área responsable.")
    add_heading(doc, "4.5 Recursos empleados", 2)
    add_table(doc,
        ["Tipo de recurso", "Descripción", "Aporte al proyecto"],
        [
            ["Información", "Archivos Excel de máquinas, áreas, localizaciones, solicitudes y repuestos.", "Permitieron estructurar el catálogo inicial y las relaciones por área."],
            ["Software", "React, Node.js, Express, MariaDB, Docker, Excel y navegador web.", "Permitieron diseñar, desarrollar, probar y ejecutar la plataforma."],
            ["Recursos humanos", "Estudiante, supervisor, técnicos, operarios y responsables de información.", "Aportaron requerimientos, validación y conocimiento del proceso."],
        ], [3.2, 7.2, 4.6])
    add_caption(doc, "Tabla 9. Recursos empleados durante el proyecto. Fuente: elaboración propia.")

    add_heading(doc, "5. Conclusiones")
    conclusions = [
        "Se diseñó e implementó INCAMAT como una solución web para centralizar el mantenimiento, las incidencias y los repuestos técnicos de Incalpaca TPX.",
        "La organización por localizaciones, áreas y máquinas facilita la consulta de activos y evita que la información quede dispersa en archivos independientes.",
        "El flujo de incidencias y órdenes de trabajo permite documentar la atención con fecha, hora, diagnóstico, evidencia fotográfica, repuestos utilizados y prueba final.",
        "La clasificación de repuestos por familia técnica, área, stock y criticidad proporciona una base para mejorar la priorización de compras y la continuidad operativa.",
        "La importación validada desde Excel y los reportes CSV permiten que la empresa adopte el sistema sin perder compatibilidad con su forma actual de trabajo.",
    ]
    for index, text in enumerate(conclusions, 1):
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(text)
        style_run(r)

    add_heading(doc, "6. Recomendaciones")
    recommendations = [
        "Validar con cada área los nombres oficiales de localización, máquina y responsable antes de efectuar nuevas cargas masivas.",
        "Realizar inventarios físicos periódicos para que el stock mostrado por el sistema corresponda al stock real de almacén.",
        "Definir responsables para evaluar la criticidad de los repuestos y revisar los valores cuando cambie la disponibilidad o el tiempo de reposición.",
        "Capacitar a operarios, técnicos, supervisores y almaceneros según el módulo que utilizarán en INCAMAT.",
        "Integrar en una etapa posterior datos de producción y calidad para calcular OEE, MTBF, MTTR y otros KPIs con información validada.",
        "Aplicar respaldos de base de datos y una política de usuarios y permisos antes de poner el sistema en producción institucional.",
    ]
    for index, text in enumerate(recommendations, 1):
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(text)
        style_run(r)

    add_heading(doc, "7. Referencias bibliográficas")
    references = [
        "Incalpaca. (s.f.). Acerca de Incalpaca. Recuperado el 13 de agosto de 2026, de https://www.incalpaca.com/es/pages/about",
        "Incalpaca TPX S.A. (2026). Organigrama Incalpaca (AGO) [Documento interno proporcionado por la empresa].",
        "TECSUP. (2026). Normas para la presentación del informe de experiencia formativa en situación real de trabajo: Módulo II [Documento de guía proporcionado].",
        "Díaz Mamani, A. N. (2026). INCAMAT: sistema web para la gestión de mantenimiento, incidencias y repuestos [Aplicación web desarrollada como proyecto de experiencia formativa].",
        "Meta. (s.f.). React: The library for web and native user interfaces. Recuperado el 13 de agosto de 2026, de https://react.dev/",
        "MariaDB Foundation. (s.f.). MariaDB Server documentation. Recuperado el 13 de agosto de 2026, de https://mariadb.com/docs/server/",
        "International Electrotechnical Commission. (2018). IEC 60812:2018: Failure modes and effects analysis (FMEA and FMECA). https://webstore.iec.ch/en/publication/26359",
        "International Organization for Standardization. (2024). ISO 55000:2024: Asset management - Vocabulary, overview and principles. https://www.iso.org/standard/83053.html",
    ]
    for reference in references:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.7)
        p.paragraph_format.first_line_indent = Cm(-0.7)
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(reference)
        style_run(r)

    new_page(doc)
    add_heading(doc, "8. Anexos")
    add_heading(doc, "Anexo 1. Hoja de tareas", 2)
    add_table(doc,
        ["Día", "Hora", "Lugar: máquina / equipo", "Descripción de tarea", "Jefe responsable"],
        [["", "", "", "", ""] for _ in range(12)],
        [1.5, 1.6, 4.0, 6.4, 3.2])
    add_caption(doc, "Tabla 10. Hoja de tareas de la experiencia formativa. Fuente: elaboración propia, basado en el formato proporcionado por TECSUP.")

    new_page(doc)
    add_heading(doc, "Anexo 2. Evidencias de INCAMAT", 2)
    add_text(doc, "Las siguientes figuras corresponden a evidencias del sistema desarrollado y de la ubicación del área de Mantenimiento dentro de la organización.")
    evidence = [
        (r"C:\Users\Admin\AppData\Local\Temp\codex-clipboard-2d60929a-4269-4103-a9d7-2b0796a4fd97.png", "Figura 1. Panel de control de INCAMAT.", "El dashboard reúne indicadores operativos, alertas de stock y el estado general de máquinas. Fuente: elaboración propia."),
        (r"C:\Users\Admin\AppData\Local\Temp\codex-clipboard-29c879e9-6790-4aa4-8e4b-7371965e670e.png", "Figura 2. Ficha técnica de máquina con código QR.", "La ficha permite consultar área, estado, marca, modelo, código interno y generar un QR para identificar el activo. Fuente: elaboración propia."),
        (r"C:\Users\Admin\AppData\Local\Temp\codex-clipboard-3d37ee7e-9351-4372-87fa-9f4420deb3d6.png", "Figura 3. Bandeja de órdenes correctivas.", "La bandeja organiza las órdenes según tipo, prioridad y estado para que el técnico inicie y atienda el trabajo pendiente. Fuente: elaboración propia."),
        (r"C:\Users\Admin\AppData\Local\Temp\codex-clipboard-41382fba-a819-4ef7-956c-4a3741e96f80.png", "Figura 4. Catálogo de repuestos por familia técnica.", "El módulo de repuestos presenta la clasificación mecánica, eléctrica y electrónica, así como filtros, criticidad y stock. Fuente: elaboración propia."),
        (r"C:\Users\Admin\AppData\Local\Temp\codex-clipboard-a3022826-c59e-472c-b580-88845fc8d714.png", "Figura 5. Validación de importación desde Excel.", "Antes de procesar un archivo, el sistema muestra filas detectadas, filas listas para importar y errores que deben corregirse. Fuente: elaboración propia."),
    ]
    for image_path, caption, explanation in evidence:
        add_figure(doc, image_path, caption, explanation)

    add_figure(
        doc,
        r"C:\IncaMant\tmp\organigrama-incalpaca.png",
        "Figura 6. Ubicación del área de Mantenimiento en el organigrama de Incalpaca.",
        "El área de Mantenimiento se encuentra bajo la Gerencia de Operaciones, junto a las plantas y áreas industriales. Fuente: Organigrama Incalpaca (AGO), documento proporcionado por la empresa.",
    )

    doc.core_properties.title = "Informe de Experiencia Formativa - INCAMAT"
    doc.core_properties.subject = "Sistema web de gestión de mantenimiento, incidencias y repuestos"
    doc.core_properties.author = "[NOMBRE DEL ESTUDIANTE]"
    enable_field_update(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    make_document()
