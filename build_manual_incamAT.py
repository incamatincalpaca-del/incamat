from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK

OUT = r"C:\\IncaMant\\Manual_de_Usuario_INCAMAT.docx"

BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
PALE_BLUE = "F4F8FC"
DARK = "0B2545"
MUTED = "5D718A"
GREEN = "2E8B57"
GOLD = "C98A00"
RED = "B9423A"

doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(1)
sec.bottom_margin = Inches(0.82)
sec.left_margin = Inches(1)
sec.right_margin = Inches(1)
sec.header_distance = Inches(.492)
sec.footer_distance = Inches(.492)

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tcMar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tcMar.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')

def set_table_widths(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tblPr = table._tbl.tblPr
    tblW = tblPr.first_child_found_in('w:tblW')
    if tblW is None:
        tblW = OxmlElement('w:tblW'); tblPr.append(tblW)
    tblW.set(qn('w:w'), '9360'); tblW.set(qn('w:type'), 'dxa')
    indent = tblPr.first_child_found_in('w:tblInd')
    if indent is None:
        indent = OxmlElement('w:tblInd'); tblPr.append(indent)
    indent.set(qn('w:w'), '120'); indent.set(qn('w:type'), 'dxa')
    grid = table._tbl.tblGrid
    for gc, width in zip(grid.gridCol_lst, widths):
        gc.set(qn('w:w'), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is None:
                tcW = OxmlElement('w:tcW'); tcPr.append(tcW)
            tcW.set(qn('w:w'), str(width)); tcW.set(qn('w:type'), 'dxa')
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    el = OxmlElement('w:tblHeader')
    el.set(qn('w:val'), 'true')
    trPr.append(el)

def set_borders(table, color='D6E0EA'):
    tblPr = table._tbl.tblPr
    borders = tblPr.first_child_found_in('w:tblBorders')
    if borders is None:
        borders = OxmlElement('w:tblBorders'); tblPr.append(borders)
    for edge in ('top','left','bottom','right','insideH','insideV'):
        tag = qn(f'w:{edge}')
        e = borders.find(tag)
        if e is None:
            e = OxmlElement(f'w:{edge}'); borders.append(e)
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), '6'); e.set(qn('w:space'), '0'); e.set(qn('w:color'), color)

def set_font(run, size=11, bold=False, color=DARK, italic=False):
    run.font.name = 'Calibri'
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)

def add_text(p, text, size=11, bold=False, color=DARK, italic=False):
    r = p.add_run(text)
    set_font(r, size, bold, color, italic)
    return r

def add_body(text, bold_lead=None):
    p = doc.add_paragraph(style='Normal')
    if bold_lead:
        add_text(p, bold_lead, bold=True)
    add_text(p, text)
    return p

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    add_text(p, text)
    return p

def add_number(text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    add_text(p, text)
    return p

def heading(text, level=1):
    p = doc.add_paragraph(style=f'Heading {level}')
    add_text(p, text, {1:16,2:13,3:12}.get(level,11), True, BLUE if level < 3 else DARK)
    return p

def add_callout(title, text, color=PALE_BLUE):
    t = doc.add_table(rows=1, cols=1)
    set_table_widths(t, [9360]); set_borders(t, 'BDD7EE')
    c = t.cell(0,0); set_cell_shading(c, color)
    p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(3)
    add_text(p, title + ' ', 11, True, BLUE)
    add_text(p, text, 10, False, DARK)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)

def add_table(headers, rows, widths):
    t=doc.add_table(rows=1, cols=len(headers))
    set_table_widths(t, widths); set_borders(t)
    h=t.rows[0]
    set_repeat_table_header(h)
    for c, text in zip(h.cells, headers):
        set_cell_shading(c, LIGHT_BLUE)
        p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.LEFT
        add_text(p, text, 9, True, BLUE)
    for row in rows:
        cells=t.add_row().cells
        for c, text in zip(cells, row):
            p=c.paragraphs[0]
            add_text(p, str(text), 9.5, False, DARK)
    doc.add_paragraph().paragraph_format.space_after=Pt(3)
    return t

def page_break():
    doc.add_page_break()

styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Calibri'; normal._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri'); normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
normal.font.size = Pt(11); normal.font.color.rgb = RGBColor.from_string(DARK)
normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.25
for name,size,before,after,color in [('Heading 1',16,18,10,BLUE),('Heading 2',13,14,7,BLUE),('Heading 3',12,10,5,DARK)]:
    s=styles[name]; s.font.name='Calibri'; s._element.rPr.rFonts.set(qn('w:ascii'),'Calibri'); s._element.rPr.rFonts.set(qn('w:hAnsi'),'Calibri'); s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(color); s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after)

# Header/footer
header=sec.header.paragraphs[0]
header.alignment=WD_ALIGN_PARAGRAPH.RIGHT
add_text(header, 'INCAMAT | Manual de usuario', 8.5, True, MUTED)
footer=sec.footer.paragraphs[0]
footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
add_text(footer, 'INCAMAT · Gestión de activos y mantenimiento · Uso interno', 8, False, MUTED)

# Cover
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(110); p.paragraph_format.space_after=Pt(8)
add_text(p,'INCAMAT',30,True,BLUE)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(20)
add_text(p,'Manual de usuario',22,True,DARK)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(38)
add_text(p,'Sistema de gestión de activos y mantenimiento\nPlanta Incalpaca TPX',13,False,MUTED)
t=doc.add_table(rows=4, cols=2); set_table_widths(t,[2700,6660]); set_borders(t,'D6E0EA')
for i,(a,b) in enumerate([('Versión','1.0'),('Elaborado para','Área de Mantenimiento – Incalpaca TPX'),('Elaborado por','Alyson Noely Diaz Mamani'),('Fecha','14 de agosto de 2026')]):
    set_cell_shading(t.cell(i,0),LIGHT_BLUE); add_text(t.cell(i,0).paragraphs[0],a,10,True,BLUE); add_text(t.cell(i,1).paragraphs[0],b,10)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(40)
add_text(p,'Documento orientado a operarios, técnicos, ingenieros y administradores del aplicativo.',10,False,MUTED,True)
page_break()

heading('Índice',1)
for item in ['1. Propósito y alcance','2. Roles y accesos','3. Ingreso al aplicativo','4. Navegación general','5. Panel de control y KPIs','6. Reportes de planta e incidencias','7. Órdenes y planificación de mantenimiento','8. Áreas, localizaciones y máquinas','9. Máquinas paradas','10. Repuestos y solicitudes','11. Importación Excel y exportación CSV','12. Usuarios y permisos','13. Recomendaciones de operación','Anexos: campos de carga y credenciales de prueba']:
    add_number(item)
add_callout('Cómo usar este manual:', 'Los procedimientos están ordenados por módulo. Siga los pasos de su perfil de usuario y complete solo los campos que correspondan a la actividad real.', 'F4F8FC')
page_break()

heading('1. Propósito y alcance',1)
add_body('INCAMAT concentra el registro de activos, incidencias, órdenes de mantenimiento, repuestos, importaciones y los indicadores operativos de la planta. El aplicativo permite que cada acción sea trazable desde el reporte inicial hasta el cierre técnico.')
heading('¿Qué puede hacer INCAMAT?',2)
for txt in ['Organizar las máquinas por área y consultar su ficha técnica con código QR.','Reportar incidencias desde planta con prioridad, descripción y evidencia fotográfica.','Atender órdenes correctivas y preventivas registrando tiempos, diagnóstico, trabajo ejecutado, repuestos y validación final.','Controlar repuestos por área, familia técnica, criticidad, stock y solicitudes.','Importar información desde Excel con validación y descargar resultados en CSV.','Visualizar indicadores de mantenimiento y actividad histórica por tipo y período.']:
    add_bullet(txt)
add_callout('Alcance del dato:', 'Los datos mostrados dependen de los registros cargados y validados en el aplicativo. Un indicador sin registros se muestra como “Sin datos” o “0”; no debe interpretarse como una medición real de desempeño.', 'FFF8E6')
heading('Estructura funcional',2)
add_table(['Origen','Proceso en INCAMAT','Resultado'],[['Operario / código QR','Reporta una incidencia con evidencia','Orden disponible para mantenimiento'],['Técnico','Inicia, diagnostica, usa repuestos y cierra','Trazabilidad técnica y estado de máquina'],['Ingeniero / administrador','Planifica, importa y revisa indicadores','Control operativo e histórico']], [2400,4000,2960])
page_break()

heading('2. Roles y accesos',1)
add_body('Los menús y acciones disponibles dependen del rol asignado. Los permisos protegen la información y separan el reporte de planta de la gestión técnica y administrativa.')
add_table(['Rol','Acceso principal','Acciones no permitidas'],[
['Operario','Reportes de planta mediante QR o formulario. Consulta básica de su reporte.','No modifica órdenes, stock, importaciones ni usuarios.'],
['Técnico','Consulta máquinas, atiende y cierra órdenes; registra repuestos usados y solicita repuestos.','No administra usuarios ni elimina importaciones.'],
['Ingeniero','Planifica mantenimientos, consulta indicadores, importa y valida información operativa.','No debe compartir credenciales ni alterar registros sin autorización.'],
['Administrador','Gestiona módulos, usuarios, permisos e importaciones, además de la operación general.','Debe controlar altas, bajas y cambios con autorización.']], [1800,4200,3360])
add_callout('Seguridad:', 'Cada usuario debe utilizar su propia cuenta. No comparta contraseñas, ni use cuentas genéricas para registrar actividades técnicas.', 'FCEEEE')
page_break()

heading('3. Ingreso al aplicativo',1)
add_number('Abra la dirección proporcionada por el administrador de INCAMAT.')
add_number('Escriba su usuario y contraseña.')
add_number('Seleccione “Iniciar sesión”. El sistema mostrará los módulos permitidos según su rol.')
add_number('Al terminar, use “Cerrar sesión”, especialmente si utiliza un equipo compartido.')
heading('Si no puede ingresar',2)
for txt in ['Verifique que el usuario y la contraseña estén escritos correctamente, respetando mayúsculas y caracteres especiales.','Solicite al administrador la creación o restablecimiento de su cuenta.','No intente acceder con cuentas de otras personas.']:
    add_bullet(txt)
add_callout('Acceso por QR:', 'El botón “Ingresar con código QR” está pensado para abrir rápidamente una ficha o un reporte desde planta. Para una incidencia, el QR identifica la máquina y evita seleccionar manualmente un equipo equivocado.', 'F4F8FC')
page_break()

heading('4. Navegación general',1)
add_body('El menú lateral organiza INCAMAT por proceso de trabajo. El nombre de cada módulo puede variar ligeramente según el rol, pero la lógica de operación se conserva.')
add_table(['Sección','Módulo','Uso'],[
['Inicio','Dashboard','Consulta rápida de KPIs, alertas e historial.'],['Operación diaria','Reportes de planta','Registro y seguimiento inicial de incidencias QR.'],['Operación diaria','Órdenes y planificación','Atención de correctivos, preventivos y agenda.'],['Activos','Áreas y localizaciones','Consulta de la estructura planta → área → máquina.'],['Activos','Máquinas / Máquinas paradas','Ficha, QR, historia y detenciones.'],['Almacén','Repuestos / Solicitudes','Stock, criticidad, movimientos y abastecimiento.'],['Control','Importaciones','Carga validada desde Excel e historial.'],['Administración','Usuarios y permisos','Alta de cuentas y definición de roles.']], [1650,2850,4860])
heading('Regla de navegación',2)
add_body('Primero ubique el área; después la máquina; finalmente la incidencia, la orden o el repuesto asociado. Este orden reduce duplicados y facilita el análisis por área.')
page_break()

heading('5. Panel de control y KPIs',1)
add_body('El Dashboard resume la situación operativa registrada en INCAMAT. Las tarjetas de historial se pueden filtrar por día, semana, mes, año o todo el historial; además, se puede alternar entre tarjetas y gráfico.')
add_table(['Indicador','Qué representa','Uso recomendado'],[
['Disponibilidad','Máquinas operativas respecto del total registrado.','Identificar impacto de paradas.'],['Cumplimiento preventivo','Mantenimientos preventivos completados respecto de los programados.','Controlar ejecución del plan.'],['Stock verificado','Proporción de inventario cuya existencia física fue confirmada.','Evitar decisiones con stock no validado.'],['MTTR','Tiempo medio registrado para resolver mantenimientos cerrados.','Medir velocidad de recuperación.'],['Máquinas paradas','Equipos detenidos por una falla, una orden pendiente o espera de repuesto.','Priorizar atención técnica.'],['Repuestos bajo mínimo','Repuestos con stock menor al mínimo configurado.','Generar o priorizar solicitudes.']], [2300,3900,3160])
add_callout('Lectura correcta:', 'Un KPI se calcula con la información disponible. Antes de comparar períodos o áreas, confirme que las órdenes, fechas, tiempos y estados estén completos.', 'FFF8E6')
page_break()

heading('6. Reportes de planta e incidencias',1)
add_body('El módulo Reportes de planta es el punto de entrada para una anomalía observada en producción. El operario puede llegar directamente escaneando el QR de la máquina.')
heading('Registrar una incidencia',2)
for txt in ['Abra “Reportes de planta” o escanee el QR de la máquina.','Seleccione primero el área. El sistema mostrará únicamente las máquinas de esa área.','Seleccione la máquina, prioridad y fecha del reporte.','Describa brevemente qué ocurre y adjunte una fotografía de cómo se encontró el equipo.','Registre el nombre de la persona que reporta y presione “Registrar falla”.']:
    add_number(txt)
heading('Prioridad',2)
add_table(['Nivel','Uso'],[['Baja','No detiene la producción; puede programarse.'],['Media','Requiere revisión técnica sin impacto crítico inmediato.'],['Alta','Afecta operación o calidad; requiere atención prioritaria.'],['Crítica','Parada, riesgo o afectación severa; debe comunicarse y atenderse con urgencia.']], [2300,7060])
add_callout('Evidencia:', 'La fotografía inicial permite comparar la condición encontrada con la condición final registrada durante el cierre de la orden.', 'F4F8FC')
page_break()

heading('7. Órdenes y planificación de mantenimiento',1)
add_body('Las incidencias recibidas y los mantenimientos planificados se administran como órdenes de trabajo. El técnico debe registrar el ciclo real de atención, no solo cambiar el estado.')
heading('Flujo de una orden correctiva',2)
add_table(['Etapa','Registro requerido','Resultado'],[
['Pendiente de iniciar','Incidencia, prioridad y evidencia inicial.','La orden queda visible en la bandeja.'],['Iniciar orden','Hora real de inicio.','Se activa el cronómetro y el estado “En atención”.'],['Diagnosticar y ejecutar','Tipo de causa, diagnóstico, trabajo realizado y repuestos usados.','Trazabilidad de la intervención.'],['Validar','Foto final y prueba final: operativa, pendiente de prueba o requiere seguimiento.','Confirmación técnica del resultado.'],['Cerrar orden','Hora final y observación de cierre.','La máquina vuelve a operativa si la prueba final lo confirma.']], [1750,4500,3110])
heading('Tipos planificados',2)
for txt in ['Preventivo planificado: tareas calendarizadas para conservar el funcionamiento del equipo.','Autónomo: checklist diario realizado por personal operativo.','Predictivo: intervención derivada de condición, medición o tendencia.','Proactivo: actividad orientada a eliminar una causa raíz o implementar una mejora.']:
    add_bullet(txt)
add_callout('Repuesto utilizado:', 'Registre el repuesto dentro de la orden. Cuando corresponde a una salida confirmada, el movimiento debe reflejarse en el control de stock.', 'FFF8E6')
page_break()

heading('8. Áreas, localizaciones y máquinas',1)
add_body('La estructura de activos organiza la información desde la planta hasta el equipo. Seleccione un área para ver únicamente las máquinas que pertenecen a ella.')
heading('Consultar una máquina',2)
for txt in ['Abra “Áreas y localizaciones” y seleccione el área, o ingrese a “Máquinas”.','Use el buscador por máquina, marca, modelo o área.','Presione una máquina para abrir su ficha.','Revise área, marca, modelo, estado, código interno, descripción breve y código QR.','Use el QR para identificación en planta y para el reporte de incidencia.']:
    add_number(txt)
add_callout('Alta manual de activos:', 'Los administradores pueden añadir un área o máquina manualmente cuando exista autorización. Complete los datos mínimos y confirme que no se trata de un registro duplicado.', 'F4F8FC')
heading('Buenas prácticas de codificación',2)
for txt in ['No modifique códigos provenientes de la fuente oficial sin autorización.','Antes de crear un activo, busque por código y nombre.','Mantenga consistente el área asignada, pues determina los equipos y repuestos visibles durante una solicitud.']:
    add_bullet(txt)
page_break()

heading('9. Máquinas paradas',1)
add_body('El módulo Máquinas paradas consolida equipos detenidos o en mantenimiento. Su propósito es responder rápidamente por qué está detenida una máquina y qué se necesita para recuperarla.')
add_table(['Dato visible','Interpretación'],[['Máquina y área','Permite asignar responsables y priorizar el impacto.'],['Motivo de parada','Falla reportada, orden pendiente o espera de repuesto.'],['Estado de atención','Pendiente, en atención, esperando repuesto o resuelta.'],['Fecha / tiempo','Ayuda a identificar antigüedad y duración de la detención.'],['Acción recomendada','Abrir la orden, solicitar repuesto o registrar avance técnico.']], [2900,6460])
add_callout('Importante:', 'Una máquina solo debe mostrarse como operativa cuando el cierre de la orden indique prueba final operativa. Si requiere seguimiento, mantenga el estado correspondiente y documente la condición.', 'FCEEEE')
page_break()

heading('10. Repuestos y solicitudes',1)
add_body('El catálogo de repuestos se clasifica por familia técnica: mecánico, eléctrico o electrónico. También permite visualizar área o unidad, criticidad, stock, consumo, fecha de solicitud y costo registrado cuando existe.')
heading('Controlar un repuesto',2)
for txt in ['Busque por código, descripción o ubicación.','Aplique filtros de área, familia técnica, stock, criticidad u ordenamiento.','Use “Evaluar” para completar el nivel de criticidad cuando el registro esté pendiente.','Use “Verificar” solo después de confirmar la existencia física de stock.','Registre movimientos de ingreso o salida para mantener trazabilidad.']:
    add_number(txt)
heading('Crear una solicitud',2)
for txt in ['Abra “Solicitudes de repuesto” y seleccione “Nueva solicitud”.','Seleccione el área solicitante; después elija una máquina opcional y el repuesto asociado a esa área.','Use el buscador para encontrar máquina, área o repuesto por texto.','Indique cantidad, prioridad, solicitante y motivo / trabajo requerido.','Envíe a aprobación. El stock no se descuenta hasta la entrega confirmada.']:
    add_number(txt)
add_callout('Criticidad:', 'La criticidad debe evaluarse con criterio técnico y de abastecimiento. Registre un nivel solo cuando se haya revisado el impacto del repuesto en la continuidad de operación.', 'FFF8E6')
page_break()

heading('11. Importación Excel y exportación CSV',1)
add_body('Las importaciones permiten actualizar información desde los archivos de trabajo de la empresa. Antes de procesar, INCAMAT muestra una vista previa y diferencia los registros a crear, actualizar o rechazar.')
heading('Proceso seguro de importación',2)
for txt in ['Abra “Importaciones” y elija el catálogo: áreas, máquinas, repuestos o mantenimientos (SRequest).','Descargue la plantilla oficial del módulo seleccionado.','Complete las columnas solicitadas sin cambiar encabezados ni códigos oficiales.','Seleccione “Elegir archivo Excel”. Revise el total de filas, errores y duplicados.','Corrija las filas observadas y procese solo cuando la validación sea correcta.','Revise el historial de importaciones para conocer archivo, fecha, tipo y estado.']:
    add_number(txt)
heading('Eliminar una importación',2)
add_body('Solo el administrador puede eliminar una importación por error. Esta acción debe borrar también los datos creados exclusivamente por ese archivo. Antes de eliminar, confirme que no haya registros posteriores vinculados que deban conservarse.')
heading('Exportar CSV',2)
add_body('Desde Órdenes y planificación use “Descargar CSV”. El archivo incluye datos de la tarea, fechas, estado, diagnóstico, trabajo u observación, repuestos utilizados y evidencia fotográfica registrada. Las órdenes completadas salen de la bandeja activa, pero se conservan en el historial y en el CSV.')
add_callout('No duplicar:', 'No vuelva a importar un archivo que ya fue procesado sin revisar su historial. El control de validación reduce duplicados, pero la revisión humana sigue siendo necesaria.', 'FCEEEE')
page_break()

heading('12. Usuarios y permisos',1)
add_body('El administrador controla las cuentas desde “Usuarios y permisos”. Debe crear usuarios individuales y asignar el rol adecuado según el trabajo que realizará cada persona.')
heading('Crear o modificar una cuenta',2)
for txt in ['Abra “Usuarios y permisos”.','Seleccione “Nuevo usuario”.','Complete usuario, nombre, correo, rol y contraseña inicial.','Guarde la cuenta y comunique la contraseña de forma segura.','Cuando cambie de puesto o salga de la empresa, modifique el rol o desactive la cuenta.']:
    add_number(txt)
add_callout('Control de acceso:', 'No use una sola cuenta compartida para un equipo de trabajo. La trazabilidad de reportes, órdenes e importaciones depende de identificar quién realizó cada acción.', 'FFF8E6')
page_break()

heading('13. Recomendaciones de operación',1)
heading('Antes de registrar',2)
for txt in ['Confirme el área y la máquina correctas.','No cree registros duplicados; primero utilice el buscador.','Tome fotografías claras, pertinentes y sin información sensible innecesaria.']:
    add_bullet(txt)
heading('Durante la atención técnica',2)
for txt in ['Inicie la orden cuando comienza realmente el trabajo.','Registre la causa, diagnóstico y trabajo realizado con lenguaje breve y claro.','Seleccione únicamente los repuestos utilizados y confirme sus cantidades.','Adjunte evidencia de cómo queda la máquina y el resultado de la prueba final.']:
    add_bullet(txt)
heading('Al cerrar',2)
for txt in ['Verifique que la hora final y la observación estén registradas.','No cierre como “operativa” una máquina que todavía requiere prueba o seguimiento.','Revise la bandeja de Máquinas paradas para confirmar que el estado se actualizó.']:
    add_bullet(txt)
heading('Soporte',2)
add_body('Para dificultades de acceso, permisos, datos maestros o importaciones, contacte al administrador de INCAMAT o al supervisor de mantenimiento. Describa el módulo, la acción realizada y, de ser posible, adjunte una captura del mensaje mostrado.')
page_break()

heading('Anexo A. Campos de carga más usados',1)
add_table(['Catálogo','Campos clave de plantilla'],[
['Áreas','código, nombre, descripción, responsable, estado.'],['Máquinas','código, máquina, área, marca, modelo, estado y descripción.'],['Repuestos','código, descripción, criticidad, stock actual, stock mínimo, área / unidad y familia técnica.'],['Mantenimientos (SRequest)','número de solicitud, fecha, hora, área, máquina, tipo, responsable, estado y descripción.']], [2500,6860])
add_callout('Plantilla oficial:', 'Utilice siempre la plantilla descargada desde INCAMAT. Las columnas pueden actualizarse según el catálogo, por lo que no se recomienda reutilizar formatos antiguos.', 'F4F8FC')
heading('Anexo B. Credenciales de prueba inicial',1)
add_body('Estas credenciales se usan únicamente para pruebas iniciales del entorno. El administrador debe cambiarlas o deshabilitarlas antes de una puesta en producción.')
add_table(['Rol','Usuario','Contraseña inicial de prueba'],[['Administrador','admin','Admin123*'],['Técnico','tecnico','Tecnico123*'],['Ingeniero','ingeniero','Ingeniero123*'],['Operario','operario','Operario123*']], [2200,3400,3760])
add_callout('Advertencia de seguridad:', 'No publique estas credenciales ni las utilice como contraseñas finales de producción.', 'FCEEEE')

doc.core_properties.title = 'Manual de usuario INCAMAT'
doc.core_properties.subject = 'Guía de uso del sistema de gestión de activos y mantenimiento'
doc.core_properties.author = 'Alyson Noely Diaz Mamani'
doc.save(OUT)
print(OUT)
